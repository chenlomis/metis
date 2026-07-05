from __future__ import annotations

import datetime as _dt
import json
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from .domain_taxonomy import render_domain_taxonomy
from .llm import complete_text


@dataclass(frozen=True)
class EvidenceUnit:
    id: str
    source_path: str
    text: str


@dataclass(frozen=True)
class TailoredArtifacts:
    output_dir: Path
    clean_resume_path: Path
    review_path: Path
    record_path: Path


_WORD_RE = re.compile(r"[a-z0-9][a-z0-9+.#-]*", re.I)

_HARD_TECH_REQUIREMENTS: list[tuple[str, tuple[str, ...]]] = [
    ("Kubernetes", ("kubernetes", "k8s", "kubecost")),
    ("containers", ("container", "containers", "docker")),
    ("multi-cluster operations", ("multi-cluster", "multicluster")),
    ("GPU utilization", ("gpu utilization", "gpu scheduling")),
    ("FinOps / cost management", ("finops", "cost management", "cost allocation")),
    ("observability", ("observability", "telemetry pipeline", "telemetry pipelines")),
    ("self-hosted infrastructure", ("self-hosted", "self hosted")),
    ("cloud-native infrastructure", ("cloud-native", "cloud native")),
]

RESUME_TAILOR_POLICY: dict[str, list[str]] = {
    "allowed": [
        "reorder emphasis by choosing the strongest relevant existing examples",
        "emphasize supported skills and experience already present in source evidence",
        "mirror JD terminology when the same concept is supported by evidence",
        "compress or modestly expand existing claims without changing their factual meaning",
        "edit qualification highlights, skills, or experience when exact source text is available",
    ],
    "forbidden": [
        "invent employers, titles, metrics, tools, models, publications, patents, or credentials",
        "convert exposure, collaboration, or adjacency into ownership",
        "claim hands-on expertise without explicit evidence",
        "add domain expertise not present in resume/profile evidence",
        "turn generic platform, cloud, API, or developer-tooling work into niche infrastructure expertise",
    ],
    "required_checks": [
        "every modified bullet maps to known evidence_ids",
        "every added keyword is grounded in evidence or appears only in the review as unsupported",
        "source_text must match the source resume before a DOCX edit is applied",
        "unsupported claims block DOCX export for that edit",
        "front matter edits are capped so qualifications/skills do not crowd out experience",
    ],
}


def _words(text: str) -> set[str]:
    stop = {
        "and", "the", "for", "with", "from", "that", "this", "into", "role",
        "team", "you", "your", "our", "their", "are", "will", "must", "have",
        "has", "pm", "product", "manager", "senior", "staff", "principal",
    }
    return {w.lower() for w in _WORD_RE.findall(text or "") if len(w) > 2 and w.lower() not in stop}


def _contains_any(text: str, patterns: tuple[str, ...]) -> bool:
    normalized = " ".join((text or "").lower().split())
    return any(pattern in normalized for pattern in patterns)


def _file_slug(value: Any) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "", str(value or ""))
    return text or "Metis"


def _normalize_resume_text(text: Any) -> str:
    normalized = " ".join(str(text or "").split())
    return re.sub(r"^[•●○▪▫*-]\s*", "", normalized).strip()


def assess_hard_technical_gap(jd_text: str, evidence: list[EvidenceUnit]) -> dict[str, Any]:
    """Detect narrow infrastructure requirements unsupported by candidate evidence."""
    jd_requirements = [
        label for label, patterns in _HARD_TECH_REQUIREMENTS
        if _contains_any(jd_text, patterns)
    ]
    if not jd_requirements:
        return {
            "is_hard_gap": False,
            "requirements": [],
            "supported": [],
            "unsupported": [],
        }

    evidence_text = "\n".join(unit.text for unit in evidence)
    supported = [
        label for label, patterns in _HARD_TECH_REQUIREMENTS
        if label in jd_requirements and _contains_any(evidence_text, patterns)
    ]
    unsupported = [label for label in jd_requirements if label not in supported]
    return {
        "is_hard_gap": len(jd_requirements) >= 4 and len(supported) <= 1,
        "requirements": jd_requirements,
        "supported": supported,
        "unsupported": unsupported,
    }


def _append_text(evidence: list[EvidenceUnit], source_path: str, value: Any) -> None:
    if not isinstance(value, str):
        return
    text = " ".join(value.split())
    if not text:
        return
    evidence.append(EvidenceUnit(f"e{len(evidence) + 1}", source_path, text))


def build_evidence_index(profile: dict) -> list[EvidenceUnit]:
    """Extract factual, resume-grounded evidence from profile.yaml.

    Preferences and aspirations are intentionally excluded from grounding. They
    are useful for scoring, but resume tailoring should only claim experience,
    skills, credentials, and demonstrated strengths.
    """
    evidence: list[EvidenceUnit] = []

    candidate = profile.get("candidate") or {}
    for idx, skill in enumerate(candidate.get("skills") or [], 1):
        _append_text(evidence, f"candidate.skills[{idx}]", skill)
    for idx, cert in enumerate(candidate.get("certifications") or [], 1):
        _append_text(evidence, f"candidate.certifications[{idx}]", cert)

    for exp_idx, exp in enumerate(profile.get("experience") or [], 1):
        if not isinstance(exp, dict):
            continue
        company = exp.get("company") or ""
        title = exp.get("title") or ""
        dates = exp.get("dates") or ""
        role_line = " ".join(p for p in [title, company, dates] if p)
        _append_text(evidence, f"experience[{exp_idx}].role", role_line)
        for h_idx, highlight in enumerate(exp.get("highlights") or [], 1):
            _append_text(evidence, f"experience[{exp_idx}].highlights[{h_idx}]", highlight)

    for edu_idx, edu in enumerate(profile.get("education") or [], 1):
        if not isinstance(edu, dict):
            continue
        line = " ".join(
            str(p)
            for p in [edu.get("degree"), edu.get("institution"), edu.get("year")]
            if p
        )
        _append_text(evidence, f"education[{edu_idx}]", line)

    for field in ("strengths", "green_flags", "yellow_flags"):
        for idx, item in enumerate(profile.get(field) or [], 1):
            _append_text(evidence, f"{field}[{idx}]", item)

    return evidence


def build_resume_evidence_index(resume_path: Path) -> list[EvidenceUnit]:
    """Extract editable evidence from the source resume DOCX paragraphs."""
    from docx import Document

    doc = Document(str(resume_path))
    evidence: list[EvidenceUnit] = []
    for idx, paragraph in enumerate(doc.paragraphs, 1):
        text = " ".join((paragraph.text or "").split())
        if len(text) < 20:
            continue
        evidence.append(EvidenceUnit(f"r{len(evidence) + 1}", f"resume.paragraphs[{idx}]", text))
    return evidence


def build_evidence_units_from_profile_index(index: dict[str, Any]) -> list[EvidenceUnit]:
    evidence: list[EvidenceUnit] = []
    for item in index.get("evidence_items") or []:
        if not isinstance(item, dict):
            continue
        claim = " ".join(str(item.get("claim") or "").split())
        evidence_id = str(item.get("id") or "").strip()
        if not claim or not evidence_id:
            continue
        source_path = str((item.get("anchors") or {}).get("profile_path") or "profile.evidence.index")
        evidence.append(EvidenceUnit(evidence_id, source_path, claim))
    for idx, item in enumerate(index.get("skill_index") or [], 1):
        if not isinstance(item, dict):
            continue
        skill = " ".join(str(item.get("skill") or "").split())
        if skill:
            evidence.append(EvidenceUnit(f"skill_{idx}", "profile.evidence.index.skill_index", skill))
    return evidence


def retrieve_evidence_for_jd(
    jd_text: str,
    evidence: list[EvidenceUnit],
    *,
    limit: int = 18,
) -> list[dict[str, Any]]:
    """Cheap deterministic retrieval before the LLM writes any edit plan."""
    jd_words = _words(jd_text)
    ranked: list[tuple[float, EvidenceUnit, list[str]]] = []
    for unit in evidence:
        unit_words = _words(unit.text)
        overlap = sorted(jd_words & unit_words)
        if not overlap:
            continue
        score = len(overlap) / max(6, len(unit_words) ** 0.5)
        ranked.append((score, unit, overlap[:12]))
    ranked.sort(key=lambda row: row[0], reverse=True)
    return [
        {
            "id": unit.id,
            "source_path": unit.source_path,
            "text": unit.text,
            "matched_terms": terms,
            "retrieval_score": round(score, 3),
        }
        for score, unit, terms in ranked[:limit]
    ]


def _extract_json(text: str) -> dict:
    raw = (text or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    start = raw.find("{")
    end = raw.rfind("}")
    if start != -1 and end != -1:
        raw = raw[start:end + 1]
    data = json.loads(raw)
    if not isinstance(data, dict):
        raise ValueError("Tailoring response must be a JSON object")
    return data


def _normalize_plan(data: dict) -> dict:
    lens = data.get("employer_lens") if isinstance(data.get("employer_lens"), dict) else {}
    candidate_lens = data.get("candidate_lens") if isinstance(data.get("candidate_lens"), dict) else {}
    edit_strategy = data.get("edit_strategy") if isinstance(data.get("edit_strategy"), dict) else {}
    edits = data.get("edits") if isinstance(data.get("edits"), list) else []
    safe_edits = []
    for edit in edits[:8]:
        if not isinstance(edit, dict):
            continue
        replacement = str(edit.get("replacement_text", "")).strip()
        if not replacement:
            continue
        safe_edits.append({
            "section": str(edit.get("section", "")).strip() or "experience",
            "target": str(edit.get("target", "")).strip(),
            "source_text": str(edit.get("source_text", "")).strip(),
            "replacement_text": replacement,
            "rationale": str(edit.get("rationale", "")).strip(),
            "evidence_ids": [str(v) for v in (edit.get("evidence_ids") or [])][:4],
            "keywords_added": [str(v) for v in (edit.get("keywords_added") or [])][:8],
        })
    return {
        "resume_tailor_policy": RESUME_TAILOR_POLICY,
        "employer_lens": {
            "fit_assessment": str(lens.get("fit_assessment", "")).strip(),
            "must_match": [str(v) for v in (lens.get("must_match") or [])][:8],
            "nice_to_have": [str(v) for v in (lens.get("nice_to_have") or [])][:8],
            "keywords": [str(v) for v in (lens.get("keywords") or [])][:32],
            "watchouts": [str(v) for v in (lens.get("watchouts") or lens.get("risks") or [])][:8],
            "knockout_concerns": [str(v) for v in (lens.get("knockout_concerns") or [])][:8],
            "seniority_scope": [str(v) for v in (lens.get("seniority_scope") or [])][:8],
            "domain_credibility_signals": [str(v) for v in (lens.get("domain_credibility_signals") or [])][:8],
            "phrase_level_language": [str(v) for v in (lens.get("phrase_level_language") or [])][:16],
        },
        "candidate_lens": {
            "explicitly_supported": [str(v) for v in (candidate_lens.get("explicitly_supported") or [])][:10],
            "adjacent_but_honest": [str(v) for v in (candidate_lens.get("adjacent_but_honest") or [])][:10],
            "unsupported_do_not_claim": [str(v) for v in (candidate_lens.get("unsupported_do_not_claim") or [])][:10],
            "safe_reframes": [str(v) for v in (candidate_lens.get("safe_reframes") or [])][:10],
        },
        "edit_strategy": {
            "primary_sections": [str(v) for v in (edit_strategy.get("primary_sections") or [])][:4],
            "max_edits": int(edit_strategy.get("max_edits") or 0),
            "rationale": str(edit_strategy.get("rationale", "")).strip(),
        },
        "summary": str(data.get("summary", "")).strip(),
        "edits": safe_edits,
        "do_not_claim": [
            str(v)
            for v in (
                data.get("do_not_claim")
                or candidate_lens.get("unsupported_do_not_claim")
                or []
            )
        ][:12],
    }


def generate_tailoring_plan(
    client: Any,
    *,
    model: str,
    role: dict,
    jd_text: str,
    profile: dict,
    evidence_matches: list[dict[str, Any]],
) -> dict:
    """Ask the LLM for a grounded, recruiter/HM-oriented edit plan."""
    candidate = (profile.get("candidate") or {}).get("name") or "Candidate"
    system = (
        "You are a recruiter, hiring manager, and ATS relevance reviewer. "
        "Your job is to tailor a resume only by reframing evidence from the source resume. "
        "Never invent employers, products, metrics, credentials, domains, tools, or scope. "
        "Prefer concise, human-readable edits over keyword stuffing. "
        "Do not lengthen a sentence unless the added words create clear recruiter or ATS signal. "
        "Use the base resume's concise professional tone."
    )
    user = json.dumps({
        "task": (
            "Identify what this JD is likely screening for, then suggest only supported "
            "resume edits grounded in the evidence_matches. Return JSON only."
        ),
        "candidate": candidate,
        "domain_transferability_reference": render_domain_taxonomy(),
        "role": {
            "title": role.get("title", ""),
            "company": role.get("company", ""),
            "location": role.get("location", ""),
            "url": role.get("url", ""),
        },
        "jd_text": jd_text[:7000],
        "evidence_matches": evidence_matches,
        "output_schema": {
            "resume_tailor_policy": {
                "allowed": RESUME_TAILOR_POLICY["allowed"],
                "forbidden": RESUME_TAILOR_POLICY["forbidden"],
                "required_checks": RESUME_TAILOR_POLICY["required_checks"],
            },
            "employer_lens": {
                "fit_assessment": "strong|moderate|weak|not_recommended, with one short reason",
                "must_match": ["specific JD requirement, not a generic PM skill"],
                "nice_to_have": ["specific JD phrase"],
                "keywords": ["ATS or HM phrase from JD"],
                "knockout_concerns": ["hard requirement likely to screen out weak candidates"],
                "seniority_scope": ["scope/seniority expectation from the JD"],
                "domain_credibility_signals": ["domain-specific proof a HM would expect"],
                "phrase_level_language": ["exact phrase from JD worth mirroring if supported"],
                "watchouts": ["important requirement not clearly supported by the source resume"],
            },
            "candidate_lens": {
                "explicitly_supported": ["JD requirement clearly supported by evidence id(s)"],
                "adjacent_but_honest": ["JD requirement partially supported but should be worded carefully"],
                "unsupported_do_not_claim": ["JD requirement not supported"],
                "safe_reframes": ["supported way to phrase candidate evidence without inflation"],
            },
            "edit_strategy": {
                "primary_sections": ["qualification highlights|skills|experience"],
                "max_edits": 0,
                "rationale": "why these sections should or should not be edited",
            },
            "summary": "two-sentence targeted resume summary, grounded in evidence",
            "edits": [{
                "section": "summary|experience|skills",
                "target": "qualification highlights, skills, company/title, or resume section",
                "source_text": "exact existing resume text to replace; do not leave blank",
                "replacement_text": "new text that preserves facts and improves JD alignment",
                "rationale": "why this helps a recruiter/HM evaluate fit",
                "evidence_ids": ["e1"],
                "keywords_added": ["keyword"],
            }],
            "do_not_claim": ["unsupported JD requirement to avoid"],
        },
        "rules": [
            "Follow resume_tailor_policy exactly. If a desired edit violates policy, put it in unsupported/do_not_claim instead of edits.",
            "First analyze the JD as a hiring manager. Separate must-haves, likely knockout concerns, seniority/scope, domain credibility signals, and exact searchable phrases.",
            "Then analyze the candidate evidence. Separate explicitly supported facts, honest adjacent evidence, unsupported claims, and safe reframes.",
            "Only after those two lenses, choose an edit strategy. Do not edit if the candidate lens shows weak or mostly adjacent support.",
            "Every edit must cite evidence_ids from evidence_matches.",
            "Every edit must use exact source_text from the source resume, including summary, qualification highlights, skills, and experience sections.",
            "If evidence is weak, put the requirement in do_not_claim instead of editing.",
            "If the JD depends on a narrow technical domain such as Kubernetes, containers, GPU utilization, FinOps, multi-cluster operations, observability, or self-hosted infrastructure and the evidence does not directly support it, mark fit_assessment as not_recommended and return no edits.",
            "Do not treat generic developer platform, API, cloud, or infrastructure language as evidence for Kubernetes/container/FinOps expertise.",
            "No inflated adjectives, no fabricated metrics, no first person.",
            "Prefer equal-length or shorter replacements; simple and concise beats broader wording.",
            "Only change text when the replacement adds explicit JD alignment not already obvious.",
            "Avoid repeating the same keyword across multiple bullets; optimize the strongest section instead of sprinkling terms everywhere.",
            "Consider qualification highlights and skills as editable sections when they are the clearest place to show supported fit.",
            "Replacement bullets should sound like a human resume, not a generated rewrite.",
            "Return 12-24 keywords/key phrases when the JD supports that many.",
            "Return at most 6 edits.",
        ],
    }, indent=2)
    response = complete_text(
        client,
        model=model,
        system=system,
        user=user,
        max_tokens=2200,
        temperature=0.2,
        json_mode=True,
    )
    return _normalize_plan(_extract_json(response.text))


def build_hard_gap_plan(role: dict, gap: dict[str, Any]) -> dict[str, Any]:
    unsupported = gap.get("unsupported") or []
    requirements = gap.get("requirements") or []
    return {
        "resume_tailor_policy": RESUME_TAILOR_POLICY,
        "employer_lens": {
            "fit_assessment": "not_recommended: JD requires narrow technical infrastructure depth not supported by the resume.",
            "must_match": requirements[:8],
            "nice_to_have": [],
            "keywords": requirements[:24],
            "knockout_concerns": requirements[:8],
            "seniority_scope": [],
            "domain_credibility_signals": requirements[:8],
            "phrase_level_language": requirements[:8],
            "watchouts": [
                "Role appears materially narrower than the candidate's platform/developer tools background.",
                "Generic platform or developer workflow evidence should not be reframed as Kubernetes/container expertise.",
            ],
        },
        "candidate_lens": {
            "explicitly_supported": gap.get("supported") or [],
            "adjacent_but_honest": ["General platform/product leadership"],
            "unsupported_do_not_claim": [f"Direct {item} expertise" for item in unsupported[:12]],
            "safe_reframes": [],
        },
        "edit_strategy": {
            "primary_sections": [],
            "max_edits": 0,
            "rationale": "Do not tailor aggressively when hard technical domain requirements are unsupported.",
        },
        "summary": "",
        "edits": [],
        "do_not_claim": [f"Direct {item} expertise" for item in unsupported[:12]],
        "technical_gap": gap,
    }


def _replace_paragraph_text(paragraph: Any, replacement: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = replacement
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replacement)


def _resume_paragraph_texts(source_resume_path: Path | None) -> set[str]:
    if not source_resume_path:
        return set()
    from docx import Document

    doc = Document(str(source_resume_path))
    return {_normalize_resume_text(paragraph.text) for paragraph in doc.paragraphs}


def _filter_supported_edits(
    edits: list[dict],
    evidence_matches: list[dict[str, Any]],
    *,
    source_resume_path: Path | None,
) -> tuple[list[dict], list[dict]]:
    known_evidence = {str(item.get("id")): item for item in evidence_matches}
    resume_texts = _resume_paragraph_texts(source_resume_path)
    supported: list[dict] = []
    rejected: list[dict] = []
    front_matter_count = 0

    for edit in edits:
        reason = ""
        evidence_ids = [str(value) for value in edit.get("evidence_ids") or []]
        source_text = _normalize_resume_text(edit.get("source_text", ""))
        replacement = " ".join(str(edit.get("replacement_text", "")).split())
        section_text = f"{edit.get('section', '')} {edit.get('target', '')}".lower()
        is_front_matter = any(term in section_text for term in ("summary", "qualification", "skills"))

        if not evidence_ids:
            reason = "missing evidence_ids"
        elif any(evidence_id not in known_evidence for evidence_id in evidence_ids):
            reason = "unknown evidence_ids"
        elif not source_text:
            reason = "missing source_text"
        elif source_resume_path and source_text not in resume_texts:
            reason = "source_text does not match source resume"
        elif len(replacement) > max(len(source_text) + 220, int(len(source_text) * 1.45)):
            reason = "replacement is too long for a concise supported edit"
        else:
            evidence_words = set()
            for evidence_id in evidence_ids:
                evidence_words |= _words(str(known_evidence[evidence_id].get("text", "")))
            replacement_words = _words(replacement)
            source_words = _words(source_text)
            if len((evidence_words | source_words) & replacement_words) < 3:
                reason = "replacement does not sufficiently overlap source evidence"
            elif is_front_matter and front_matter_count >= 2:
                reason = "front matter edit limit exceeded"

        if reason:
            rejected.append({**edit, "rejection_reason": reason})
        else:
            if is_front_matter:
                front_matter_count += 1
            supported.append(edit)

    return supported, rejected


def _write_updated_resume_docx(source_resume_path: Path, output_path: Path, edits: list[dict]) -> list[dict]:
    from docx import Document

    shutil.copy2(source_resume_path, output_path)
    doc = Document(str(output_path))
    applied: list[dict] = []
    for edit in edits:
        source_text = _normalize_resume_text(edit.get("source_text", ""))
        replacement = " ".join(str(edit.get("replacement_text", "")).split())
        if not source_text or not replacement:
            continue
        for paragraph in doc.paragraphs:
            original = " ".join((paragraph.text or "").split())
            if _normalize_resume_text(original) == source_text:
                if re.match(r"^[•●○▪▫*-]\s+", original) and not re.match(r"^[•●○▪▫*-]\s+", replacement):
                    replacement = f"{original.split()[0]} {replacement}"
                _replace_paragraph_text(paragraph, replacement)
                applied.append({
                    "source_text": source_text,
                    "replacement_text": replacement,
                    "target": edit.get("target", ""),
                })
                break
    doc.save(str(output_path))
    output_path.chmod(0o600)
    return applied


def _render_review_markdown(role: dict, plan: dict, evidence_matches: list[dict]) -> str:
    lines = [
        f"# Tailoring Review: {role.get('title', 'Role')} at {role.get('company', 'Company')}",
        "",
        f"Generated: {_dt.datetime.now().isoformat(timespec='seconds')}",
    ]
    if role.get("url"):
        lines.append(f"Posting: {role['url']}")

    policy = plan.get("resume_tailor_policy") or RESUME_TAILOR_POLICY
    if policy:
        lines += ["", "## Resume Tailoring Policy"]
        for key, label in [
            ("allowed", "Allowed"),
            ("forbidden", "Forbidden"),
            ("required_checks", "Required Checks"),
        ]:
            values = policy.get(key) or []
            if values:
                lines.append(f"### {label}")
                lines += [f"- {value}" for value in values]

    lens = plan.get("employer_lens") or {}
    if lens.get("fit_assessment"):
        lines += ["", "## Fit Assessment", f"- {lens['fit_assessment']}"]
    for key, label in [
        ("must_match", "Must Match"),
        ("nice_to_have", "Nice To Have"),
        ("knockout_concerns", "Likely Knockout Concerns"),
        ("seniority_scope", "Seniority / Scope"),
        ("domain_credibility_signals", "Domain Credibility Signals"),
        ("phrase_level_language", "Phrase-Level Language"),
        ("keywords", "Keywords"),
        ("watchouts", "Watchouts / Unsupported Requirements"),
    ]:
        values = lens.get(key) or []
        if values:
            lines += ["", f"## {label}"]
            lines += [f"- {value}" for value in values]

    candidate_lens = plan.get("candidate_lens") or {}
    for key, label in [
        ("explicitly_supported", "Explicitly Supported"),
        ("adjacent_but_honest", "Adjacent But Honest"),
        ("unsupported_do_not_claim", "Unsupported / Do Not Claim"),
        ("safe_reframes", "Safe Reframes"),
    ]:
        values = candidate_lens.get(key) or []
        if values:
            lines += ["", f"## Candidate Evidence Lens: {label}"]
            lines += [f"- {value}" for value in values]

    edit_strategy = plan.get("edit_strategy") or {}
    if edit_strategy:
        lines += ["", "## Edit Strategy"]
        if edit_strategy.get("primary_sections"):
            lines.append(f"- Primary sections: {', '.join(edit_strategy['primary_sections'])}")
        if edit_strategy.get("max_edits") is not None:
            lines.append(f"- Max edits: {edit_strategy.get('max_edits')}")
        if edit_strategy.get("rationale"):
            lines.append(f"- Why: {edit_strategy['rationale']}")

    if plan.get("edits"):
        lines += ["", "## Proposed Changes"]
        for idx, edit in enumerate(plan["edits"], 1):
            lines += [
                "",
                f"### Change {idx}: {edit.get('target') or edit.get('section')}",
            ]
            if edit.get("source_text"):
                lines.append(f"- Before: {edit['source_text']}")
            lines.append(f"- After: {edit.get('replacement_text', '')}")
            if edit.get("rationale"):
                lines.append(f"- Why: {edit['rationale']}")
            if edit.get("evidence_ids"):
                lines.append(f"- Evidence: {', '.join(edit['evidence_ids'])}")
            if edit.get("keywords_added"):
                lines.append(f"- Keywords: {', '.join(edit['keywords_added'])}")

    if plan.get("unsupported_edits"):
        lines += ["", "## Unsupported / Not Applied"]
        for idx, edit in enumerate(plan["unsupported_edits"], 1):
            lines += [
                "",
                f"### Unsupported {idx}: {edit.get('target') or edit.get('section')}",
                f"- Reason: {edit.get('rejection_reason', 'unsupported')}",
            ]
            if edit.get("source_text"):
                lines.append(f"- Before: {edit['source_text']}")
            if edit.get("replacement_text"):
                lines.append(f"- Proposed: {edit['replacement_text']}")
            if edit.get("evidence_ids"):
                lines.append(f"- Evidence: {', '.join(edit['evidence_ids'])}")

    if plan.get("do_not_claim"):
        lines += ["", "## Do Not Claim"]
        lines += [f"- {item}" for item in plan["do_not_claim"]]

    if evidence_matches:
        lines += ["", "## Evidence Used"]
        for evidence in evidence_matches:
            lines.append(f"- {evidence['id']} ({evidence['source_path']}): {evidence['text']}")

    return "\n".join(lines).strip() + "\n"


def write_tailoring_artifacts(
    *,
    role: dict,
    profile: dict,
    jd_text: str,
    evidence_matches: list[dict[str, Any]],
    plan: dict,
    output_dir: Path,
    source_resume_path: Path | None = None,
) -> TailoredArtifacts:
    """Write updated resume DOCX, review markdown, and audit record for one role."""
    output_dir.mkdir(mode=0o700, parents=True, exist_ok=True)
    candidate_name = _file_slug((profile.get("candidate") or {}).get("name") or "Candidate")
    company_name = _file_slug(role.get("company") or "Company")
    clean_path = output_dir / f"{candidate_name}_{company_name}_resume.docx" if source_resume_path else output_dir / f"{candidate_name}_{company_name}_resume.md"
    review_path = output_dir / "resume_review.md"
    record_path = output_dir / "tailoring_record.json"
    applied_edits: list[dict] = []
    supported_edits, unsupported_edits = _filter_supported_edits(
        plan.get("edits") or [],
        evidence_matches,
        source_resume_path=source_resume_path,
    )
    review_plan = {
        **plan,
        "edits": supported_edits,
        "unsupported_edits": unsupported_edits,
    }

    if source_resume_path:
        applied_edits = _write_updated_resume_docx(source_resume_path, clean_path, supported_edits)
    else:
        clean_path.write_text(
            "# Updated resume unavailable\n\nProvide a source resume DOCX so Metis can create an updated resume copy.\n",
            encoding="utf-8",
        )
    review_path.write_text(_render_review_markdown(role, review_plan, evidence_matches), encoding="utf-8")
    record = {
        "created_at": _dt.datetime.now().isoformat(timespec="seconds"),
        "role": role,
        "jd_excerpt": jd_text[:1200],
        "evidence_matches": evidence_matches,
        "plan": review_plan,
        "applied_edits": applied_edits,
        "unsupported_edits": unsupported_edits,
        "artifacts": {
            "clean_resume": str(clean_path),
            "review": str(review_path),
        },
    }
    record_path.write_text(json.dumps(record, indent=2), encoding="utf-8")
    for path in (clean_path, review_path, record_path):
        path.chmod(0o600)
    return TailoredArtifacts(output_dir, clean_path, review_path, record_path)
