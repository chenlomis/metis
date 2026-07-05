from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock

from docx import Document


def _profile() -> dict:
    return {
        "candidate": {
            "name": "Lomis Chen",
            "email": "chenlomis@gmail.com",
            "location": "Redmond, WA",
            "skills": ["Developer platforms", "AI infrastructure", "Enterprise SaaS"],
        },
        "aspirations": {
            "company_types": ["ai-infrastructure"],
        },
        "preferences": {
            "company_stage": ["growth-stage"],
        },
        "experience": [
            {
                "company": "DocuSign",
                "title": "Senior Product Manager",
                "dates": "2021-2025",
                "highlights": [
                    "Led AI-powered developer workflow improvements for enterprise platform teams.",
                    "Partnered with engineering to ship API and deployment capabilities.",
                ],
            }
        ],
        "strengths": ["Cross-functional product leadership with engineering-heavy teams."],
        "education": [{"degree": "MBA", "institution": "University of Washington"}],
    }


def _resume_docx(tmp_path: Path) -> Path:
    path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Lomis Chen")
    doc.add_paragraph("Partnered with engineering to ship API and deployment capabilities.")
    doc.add_paragraph("Led AI-powered developer workflow improvements for enterprise platform teams.")
    doc.save(path)
    return path


def test_build_evidence_index_uses_facts_not_preferences():
    from metis.tailor import build_evidence_index

    evidence = build_evidence_index(_profile())
    texts = [unit.text for unit in evidence]

    assert "Developer platforms" in texts
    assert "Led AI-powered developer workflow improvements for enterprise platform teams." in texts
    assert all("growth-stage" not in text for text in texts)
    assert all("ai-infrastructure" not in text for text in texts)


def test_retrieve_evidence_for_jd_ranks_relevant_resume_facts():
    from metis.tailor import build_evidence_index, retrieve_evidence_for_jd

    evidence = build_evidence_index(_profile())
    matches = retrieve_evidence_for_jd(
        "This role owns developer platform APIs, deployment workflows, and engineering velocity.",
        evidence,
        limit=3,
    )

    assert matches
    assert any("developer workflow" in match["text"] for match in matches)
    assert any("matched_terms" in match for match in matches)


def test_hard_technical_gap_detects_unsupported_container_depth():
    from metis.tailor import assess_hard_technical_gap, build_evidence_index

    evidence = build_evidence_index(_profile())
    jd = (
        "Own Kubernetes-native containers strategy for Kubecost, GPU utilization, "
        "multi-cluster operations, observability, telemetry pipelines, FinOps cost allocation, "
        "and self-hosted infrastructure."
    )

    gap = assess_hard_technical_gap(jd, evidence)

    assert gap["is_hard_gap"] is True
    assert "Kubernetes" in gap["unsupported"]
    assert "FinOps / cost management" in gap["unsupported"]


def test_normalize_plan_preserves_employer_and_candidate_lenses():
    from metis.tailor import _normalize_plan

    plan = _normalize_plan({
        "employer_lens": {
            "fit_assessment": "moderate: close platform fit",
            "must_match": ["developer workflows"],
            "knockout_concerns": ["security domain"],
            "seniority_scope": ["staff-level roadmap ownership"],
            "domain_credibility_signals": ["developer-facing APIs"],
            "phrase_level_language": ["developer workflows"],
        },
        "candidate_lens": {
            "explicitly_supported": ["Azure CLI developer-facing APIs"],
            "adjacent_but_honest": ["enterprise platform adoption"],
            "unsupported_do_not_claim": ["identity security ownership"],
            "safe_reframes": ["developer tooling"],
        },
        "edit_strategy": {
            "primary_sections": ["qualification highlights", "experience"],
            "max_edits": 3,
            "rationale": "Supported fit appears in highlights and Azure CLI experience.",
        },
    })

    assert plan["employer_lens"]["knockout_concerns"] == ["security domain"]
    assert plan["candidate_lens"]["unsupported_do_not_claim"] == ["identity security ownership"]
    assert plan["edit_strategy"]["primary_sections"] == ["qualification highlights", "experience"]
    assert "allowed" in plan["resume_tailor_policy"]


def test_hard_gap_plan_returns_no_edits():
    from metis.tailor import build_hard_gap_plan

    plan = build_hard_gap_plan(
        {"title": "Principal Product Manager - Containers", "company": "IBM"},
        {
            "is_hard_gap": True,
            "requirements": ["Kubernetes", "containers", "observability", "FinOps / cost management"],
            "supported": [],
            "unsupported": ["Kubernetes", "containers"],
        },
    )

    assert plan["employer_lens"]["fit_assessment"].startswith("not_recommended")
    assert plan["candidate_lens"]["unsupported_do_not_claim"]
    assert plan["edit_strategy"]["max_edits"] == 0
    assert "forbidden" in plan["resume_tailor_policy"]
    assert plan["edits"] == []
    assert "Direct Kubernetes expertise" in plan["do_not_claim"]


def test_build_resume_evidence_index_reads_source_docx(tmp_path):
    from metis.tailor import build_resume_evidence_index

    evidence = build_resume_evidence_index(_resume_docx(tmp_path))

    assert evidence[0].id == "r1"
    assert evidence[0].source_path.startswith("resume.paragraphs")
    assert "API and deployment capabilities" in evidence[0].text


def test_write_tailoring_artifacts_creates_updated_docx_review_and_record(tmp_path):
    from metis.tailor import write_tailoring_artifacts

    plan = {
        "employer_lens": {
            "fit_assessment": "strong: developer platform match",
            "must_match": ["developer platform"],
            "keywords": ["API"],
            "knockout_concerns": ["developer-facing product ownership"],
            "domain_credibility_signals": ["developer platform APIs"],
        },
        "candidate_lens": {
            "explicitly_supported": ["API delivery"],
            "adjacent_but_honest": ["enterprise platform adoption"],
            "unsupported_do_not_claim": ["Kubernetes ownership"],
            "safe_reframes": ["developer platform API delivery"],
        },
        "edit_strategy": {
            "primary_sections": ["experience"],
            "max_edits": 1,
            "rationale": "Experience contains the strongest supported API evidence.",
        },
        "summary": "Product leader focused on developer platforms and enterprise AI workflows.",
        "edits": [{
            "section": "experience",
            "target": "DocuSign",
            "source_text": "Partnered with engineering to ship API and deployment capabilities.",
            "replacement_text": "Partnered with engineering to ship developer platform API and deployment capabilities for enterprise teams.",
            "rationale": "Surfaces JD language while preserving the original fact.",
            "evidence_ids": ["e3"],
            "keywords_added": ["developer platform", "API"],
        }],
        "do_not_claim": ["Kubernetes ownership"],
    }
    resume_path = _resume_docx(tmp_path)

    artifacts = write_tailoring_artifacts(
        role={"title": "Staff PM", "company": "Acme", "url": "https://example.com/job"},
        profile=_profile(),
        jd_text="developer platform API deployment",
        evidence_matches=[{
            "id": "e3",
            "source_path": "resume.paragraphs[2]",
            "text": "Partnered with engineering to ship API and deployment capabilities.",
        }],
        plan=plan,
        output_dir=tmp_path / "tailored",
        source_resume_path=resume_path,
    )

    clean_doc = Document(str(artifacts.clean_resume_path))
    clean = "\n".join(p.text for p in clean_doc.paragraphs)
    review = artifacts.review_path.read_text()
    record = json.loads(artifacts.record_path.read_text())

    assert "developer platform API and deployment capabilities" in clean
    assert artifacts.clean_resume_path.name == "LomisChen_Acme_resume.docx"
    assert not (artifacts.output_dir / "resume_tailored.md").exists()
    assert "## Resume Tailoring Policy" in review
    assert "## Fit Assessment" in review
    assert "## Candidate Evidence Lens: Explicitly Supported" in review
    assert "## Edit Strategy" in review
    assert "Before:" in review
    assert record["plan"]["do_not_claim"] == ["Kubernetes ownership"]
    assert record["applied_edits"][0]["target"] == "DocuSign"


def test_write_tailoring_artifacts_rejects_unknown_evidence_ids(tmp_path):
    from metis.tailor import write_tailoring_artifacts

    plan = {
        "employer_lens": {"must_match": ["developer platform"]},
        "edits": [{
            "section": "experience",
            "target": "DocuSign",
            "source_text": "Partnered with engineering to ship API and deployment capabilities.",
            "replacement_text": "Partnered with engineering to ship Kubernetes platform APIs for enterprise teams.",
            "rationale": "Unsupported keyword stuffing.",
            "evidence_ids": ["missing"],
            "keywords_added": ["Kubernetes"],
        }],
        "do_not_claim": [],
    }
    resume_path = _resume_docx(tmp_path)

    artifacts = write_tailoring_artifacts(
        role={"title": "Staff PM", "company": "Acme"},
        profile=_profile(),
        jd_text="developer platform API deployment",
        evidence_matches=[],
        plan=plan,
        output_dir=tmp_path / "tailored",
        source_resume_path=resume_path,
    )

    clean_doc = Document(str(artifacts.clean_resume_path))
    clean = "\n".join(p.text for p in clean_doc.paragraphs)
    record = json.loads(artifacts.record_path.read_text())

    assert "Kubernetes platform APIs" not in clean
    assert record["applied_edits"] == []
    assert record["unsupported_edits"][0]["rejection_reason"] == "unknown evidence_ids"


def test_write_tailoring_artifacts_matches_bulleted_source_text_without_glyph(tmp_path):
    from metis.tailor import write_tailoring_artifacts

    path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("• Proven ability to translate product vision to delivery.")
    doc.save(path)
    plan = {
        "employer_lens": {"must_match": ["developer workflows"]},
        "edits": [{
            "section": "qualification highlights",
            "target": "Qualifications",
            "source_text": "Proven ability to translate product vision to delivery.",
            "replacement_text": "Proven ability to translate product vision into developer workflow improvements.",
            "rationale": "Clarifies supported developer workflow alignment.",
            "evidence_ids": ["r1"],
            "keywords_added": ["developer workflow"],
        }],
        "do_not_claim": [],
    }

    artifacts = write_tailoring_artifacts(
        role={"title": "Staff PM", "company": "Acme"},
        profile=_profile(),
        jd_text="developer workflows",
        evidence_matches=[{
            "id": "r1",
            "source_path": "resume.paragraphs[1]",
            "text": "• Proven ability to translate product vision to delivery.",
        }],
        plan=plan,
        output_dir=tmp_path / "tailored",
        source_resume_path=path,
    )

    clean_doc = Document(str(artifacts.clean_resume_path))
    assert clean_doc.paragraphs[0].text.startswith("• Proven ability")
    assert "developer workflow improvements" in clean_doc.paragraphs[0].text


def test_write_tailoring_artifacts_caps_front_matter_edits(tmp_path):
    from metis.tailor import write_tailoring_artifacts

    path = tmp_path / "resume.docx"
    doc = Document()
    for idx in range(3):
        doc.add_paragraph(f"• Qualification highlight {idx} about developer platforms.")
    doc.save(path)
    edits = [
        {
            "section": "qualification highlights",
            "target": "Qualifications",
            "source_text": f"Qualification highlight {idx} about developer platforms.",
            "replacement_text": f"Qualification highlight {idx} about developer workflow platforms.",
            "rationale": "Clarifies supported developer workflow alignment.",
            "evidence_ids": [f"r{idx + 1}"],
            "keywords_added": ["developer workflow"],
        }
        for idx in range(3)
    ]

    artifacts = write_tailoring_artifacts(
        role={"title": "Staff PM", "company": "Acme"},
        profile=_profile(),
        jd_text="developer workflows",
        evidence_matches=[
            {
                "id": f"r{idx + 1}",
                "source_path": f"resume.paragraphs[{idx + 1}]",
                "text": f"• Qualification highlight {idx} about developer platforms.",
            }
            for idx in range(3)
        ],
        plan={"employer_lens": {}, "edits": edits, "do_not_claim": []},
        output_dir=tmp_path / "tailored",
        source_resume_path=path,
    )
    record = json.loads(artifacts.record_path.read_text())

    assert len(record["applied_edits"]) == 2
    assert record["unsupported_edits"][0]["rejection_reason"] == "front matter edit limit exceeded"


def test_run_resume_tailor_with_recent_role_uses_mocked_fetch_and_llm(tmp_path, monkeypatch):
    import metis.resume_cmd as resume_cmd

    monkeypatch.setenv("METIS_DATA_DIR", str(tmp_path))
    roles = [{
        "title": "Staff Product Manager",
        "company": "Acme",
        "url": "https://www.linkedin.com/jobs/view/123/",
        "eval": {"verdict": "apply", "score": 81},
    }]
    monkeypatch.setattr(resume_cmd, "load_profile_yaml", lambda: _profile())
    monkeypatch.setattr(resume_cmd, "_load_recent_roles", lambda limit=40: roles)
    monkeypatch.setattr(resume_cmd, "_resolve_resume_path", lambda profile, resume_path=None: _resume_docx(tmp_path))
    monkeypatch.setattr(resume_cmd, "_fetch_jd", lambda role: "developer platform API deployment")
    monkeypatch.setattr(resume_cmd, "create_llm_client", lambda provider, api_key: MagicMock())
    monkeypatch.setattr(
        resume_cmd,
        "generate_tailoring_plan",
        lambda *args, **kwargs: {
            "employer_lens": {"must_match": ["developer platform"], "keywords": ["API"]},
            "summary": "Product leader focused on developer platforms.",
            "edits": [],
            "do_not_claim": [],
        },
    )

    artifacts = resume_cmd.run_resume_tailor(
        api_key="test-key",
        out_dir=str(tmp_path),
        tailor_all=True,
    )

    assert len(artifacts) == 1
    assert Path(artifacts[0]["clean_resume"]).exists()
    assert artifacts[0]["clean_resume"].endswith("LomisChen_Acme_resume.docx")
    assert "Staff Product Manager at Acme" in artifacts[0]["role"]


def test_non_interactive_without_all_requires_explicit_batch_choice(monkeypatch):
    import pytest
    import metis.resume_cmd as resume_cmd

    monkeypatch.setattr(resume_cmd, "load_profile_yaml", lambda: _profile())
    monkeypatch.setattr(resume_cmd, "_resolve_resume_path", lambda profile, resume_path=None: Path("/tmp/resume.docx"))
    monkeypatch.setattr(
        resume_cmd,
        "_load_recent_roles",
        lambda limit=40: [{
            "title": "Staff PM",
            "company": "Acme",
            "url": "https://example.com/1",
            "eval": {"verdict": "apply", "score": 81},
        }],
    )

    with pytest.raises(SystemExit) as exc:
        resume_cmd.run_resume_tailor(api_key="test-key", non_interactive=True)

    assert "resume tailor --all" in str(exc.value)


def test_run_resume_tailor_rejects_output_outside_data_dir(tmp_path, monkeypatch):
    import pytest
    import metis.resume_cmd as resume_cmd

    data_dir = tmp_path / "data"
    outside = tmp_path / "outside"
    monkeypatch.setenv("METIS_DATA_DIR", str(data_dir))
    monkeypatch.setattr(resume_cmd, "load_profile_yaml", lambda: _profile())
    monkeypatch.setattr(resume_cmd, "_resolve_resume_path", lambda profile, resume_path=None: _resume_docx(tmp_path))
    monkeypatch.setattr(
        resume_cmd,
        "_load_recent_roles",
        lambda limit=40: [{
            "title": "Staff PM",
            "company": "Acme",
            "url": "https://example.com/1",
            "eval": {"verdict": "apply", "score": 81},
        }],
    )
    monkeypatch.setattr(resume_cmd, "create_llm_client", lambda provider, api_key: MagicMock())

    with pytest.raises(SystemExit) as exc:
        resume_cmd.run_resume_tailor(
            api_key="test-key",
            out_dir=str(outside),
            tailor_all=True,
        )

    assert "METIS_DATA_DIR" in str(exc.value)


def test_recent_roles_without_url_fail_before_picker(monkeypatch):
    import pytest
    import metis.resume_cmd as resume_cmd

    monkeypatch.setattr(resume_cmd, "load_profile_yaml", lambda: _profile())
    monkeypatch.setattr(resume_cmd, "_resolve_resume_path", lambda profile, resume_path=None: Path("/tmp/resume.docx"))
    monkeypatch.setattr(
        resume_cmd,
        "_load_recent_roles",
        lambda limit=40: [{
            "title": "Staff PM",
            "company": "Acme",
            "eval": {"verdict": "apply", "score": 80},
        }],
    )

    with pytest.raises(SystemExit) as exc:
        resume_cmd.run_resume_tailor(api_key="test-key", non_interactive=True)

    assert "do not include posting URLs yet" in str(exc.value)


def test_fetch_jd_prefers_stored_jd_without_network(monkeypatch):
    import metis.resume_cmd as resume_cmd

    def fail_fetch(*args, **kwargs):
        raise AssertionError("network fetch should not run")

    monkeypatch.setattr(resume_cmd, "_validate_fetch_url", fail_fetch)

    assert resume_cmd._fetch_jd({"url": "http://127.0.0.1/job", "jd": "stored JD"}) == "stored JD"


def test_validate_fetch_url_blocks_loopback(monkeypatch):
    import socket
    import pytest
    import metis.resume_cmd as resume_cmd

    monkeypatch.setattr(
        socket,
        "getaddrinfo",
        lambda *args, **kwargs: [(socket.AF_INET, socket.SOCK_STREAM, 0, "", ("127.0.0.1", 443))],
    )

    with pytest.raises(ValueError):
        resume_cmd._validate_fetch_url("https://example.com/job")


def test_validate_fetch_url_rejects_non_http_scheme():
    import pytest
    import metis.resume_cmd as resume_cmd

    with pytest.raises(ValueError):
        resume_cmd._validate_fetch_url("file:///etc/passwd")


def test_picker_roles_sort_by_score_descending():
    from metis.resume_cmd import _sort_roles_for_picker

    roles = [
        {"title": "Low", "company": "Acme", "eval": {"score": 55}},
        {"title": "High", "company": "Beta", "eval": {"score": 84}},
        {"title": "Mid", "company": "Core", "eval": {"score": 70}},
    ]

    assert [role["title"] for role in _sort_roles_for_picker(roles)] == ["High", "Mid", "Low"]


def test_output_dir_uses_short_readable_role_folder(tmp_path):
    from metis.resume_cmd import _make_output_dir

    path = _make_output_dir(
        tmp_path,
        {
            "company": "IBM",
            "title": "Principal Product Manager - Containers",
            "role_hash": "abc123",
        },
    )

    assert path.parent.name.isdigit()
    assert path.name == "ibm_containers"


def test_select_all_expands_to_every_sorted_role():
    from metis.resume_cmd import _SELECT_ALL, _expand_selection

    roles = [
        {"title": "High", "company": "Beta", "eval": {"score": 84}},
        {"title": "Mid", "company": "Core", "eval": {"score": 70}},
    ]

    assert _expand_selection([_SELECT_ALL], roles) == roles


def test_expand_selection_uses_stable_role_ids():
    from metis.resume_cmd import _expand_selection

    roles = [
        {"title": "High", "company": "Beta", "role_hash": "r1", "eval": {"score": 84}},
        {"title": "Mid", "company": "Core", "role_hash": "r2", "eval": {"score": 70}},
    ]

    assert _expand_selection(["r2"], roles) == [roles[1]]


def test_cancel_selection_exits_cleanly():
    import pytest
    from metis.resume_cmd import _CANCEL, _expand_selection

    with pytest.raises(SystemExit) as exc:
        _expand_selection([_CANCEL], [{"title": "Role"}])

    assert "Cancelled" in str(exc.value)


def test_load_recent_roles_keeps_latest_trace_date(tmp_path, monkeypatch):
    import json
    import metis.resume_cmd as resume_cmd

    runs = tmp_path / "runs.jsonl"
    rows = [
        {
            "ts": "2026-07-03T18:34:09",
            "title": "Old",
            "company": "Snorkel AI",
            "url": "https://example.com/old",
            "eval": {"verdict": "consider", "score": 70},
        },
        {
            "ts": "2026-07-04T10:34:03",
            "title": "New",
            "company": "Jobgether",
            "url": "https://example.com/new",
            "eval": {"verdict": "apply", "score": 82},
        },
    ]
    runs.write_text("\n".join(json.dumps(row) for row in rows), encoding="utf-8")
    monkeypatch.setattr(resume_cmd, "RUNS_PATH", runs)

    loaded = resume_cmd._load_recent_roles()

    assert [role["title"] for role in loaded] == ["New"]


def test_not_recommended_tailoring_status_hides_role(tmp_path):
    import json
    from metis.resume_cmd import _eligible_recent_roles, _tailoring_status_by_role_hash

    record_dir = tmp_path / "resume_tailor" / "20260704" / "ibm_containers"
    record_dir.mkdir(parents=True)
    (record_dir / "tailoring_record.json").write_text(json.dumps({
        "role": {"title": "Principal Product Manager - Containers", "company": "IBM", "role_hash": "ibm1"},
        "plan": {"employer_lens": {"fit_assessment": "not_recommended: hard gap"}},
        "artifacts": {"clean_resume": "x"},
    }), encoding="utf-8")
    older_dir = tmp_path / "resume_tailor" / "20260704" / "ibm_old"
    older_dir.mkdir(parents=True)
    (older_dir / "tailoring_record.json").write_text(json.dumps({
        "role": {"title": "Principal Product Manager - Containers", "company": "IBM", "role_hash": "ibm1"},
        "plan": {"employer_lens": {"fit_assessment": "strong"}},
        "artifacts": {"clean_resume": "x"},
    }), encoding="utf-8")
    statuses = _tailoring_status_by_role_hash(tmp_path)

    roles = [
        {"title": "Principal Product Manager - Containers", "company": "IBM", "role_hash": "ibm1", "url": "https://example.com/ibm"},
        {"title": "Staff PM", "company": "Jobgether", "role_hash": "job1", "url": "https://example.com/job"},
    ]

    assert statuses["ibm1"] == "not_recommended"
    assert [role["company"] for role in _eligible_recent_roles(roles, statuses)] == ["Jobgether"]


def test_tailor_all_uses_every_eligible_recent_role(tmp_path, monkeypatch):
    import metis.resume_cmd as resume_cmd

    monkeypatch.setenv("METIS_DATA_DIR", str(tmp_path))
    roles = [
        {"title": "Staff PM", "company": "Acme", "url": "https://example.com/1", "eval": {"verdict": "apply", "score": 81}},
        {"title": "Lead PM", "company": "Beta", "url": "https://example.com/2", "eval": {"verdict": "consider", "score": 62}},
    ]
    monkeypatch.setattr(resume_cmd, "load_profile_yaml", lambda: _profile())
    monkeypatch.setattr(resume_cmd, "_load_recent_roles", lambda limit=40: roles)
    monkeypatch.setattr(resume_cmd, "_resolve_resume_path", lambda profile, resume_path=None: _resume_docx(tmp_path))
    monkeypatch.setattr(resume_cmd, "_fetch_jd", lambda role: "developer platform API deployment")
    monkeypatch.setattr(resume_cmd, "create_llm_client", lambda provider, api_key: MagicMock())
    monkeypatch.setattr(
        resume_cmd,
        "generate_tailoring_plan",
        lambda *args, **kwargs: {
            "employer_lens": {"must_match": ["developer platform"], "keywords": ["API"]},
            "summary": "Product leader focused on developer platforms.",
            "edits": [],
            "do_not_claim": [],
        },
    )

    artifacts = resume_cmd.run_resume_tailor(
        api_key="test-key",
        out_dir=str(tmp_path),
        tailor_all=True,
    )

    assert [item["role"] for item in artifacts] == ["Staff PM at Acme", "Lead PM at Beta"]


def test_top_n_uses_highest_score_eligible_roles(tmp_path, monkeypatch):
    import metis.resume_cmd as resume_cmd

    monkeypatch.setenv("METIS_DATA_DIR", str(tmp_path))
    roles = [
        {"title": "Mid PM", "company": "Beta", "url": "https://example.com/2", "eval": {"verdict": "consider", "score": 62}},
        {"title": "Staff PM", "company": "Acme", "url": "https://example.com/1", "eval": {"verdict": "apply", "score": 81}},
        {"title": "Principal PM", "company": "Core", "url": "https://example.com/3", "eval": {"verdict": "apply", "score": 90}},
    ]
    monkeypatch.setattr(resume_cmd, "load_profile_yaml", lambda: _profile())
    monkeypatch.setattr(resume_cmd, "_load_recent_roles", lambda limit=40: roles)
    monkeypatch.setattr(resume_cmd, "_resolve_resume_path", lambda profile, resume_path=None: _resume_docx(tmp_path))
    monkeypatch.setattr(resume_cmd, "_fetch_jd", lambda role: "developer platform API deployment")
    monkeypatch.setattr(resume_cmd, "create_llm_client", lambda provider, api_key: MagicMock())
    monkeypatch.setattr(
        resume_cmd,
        "generate_tailoring_plan",
        lambda *args, **kwargs: {
            "employer_lens": {"must_match": ["developer platform"], "keywords": ["API"]},
            "summary": "Product leader focused on developer platforms.",
            "edits": [],
            "do_not_claim": [],
        },
    )

    artifacts = resume_cmd.run_resume_tailor(
        api_key="test-key",
        top_n=2,
    )

    assert [item["role"] for item in artifacts] == ["Principal PM at Core", "Staff PM at Acme"]
